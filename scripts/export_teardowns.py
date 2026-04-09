import os
import glob
import argparse
from docx import Document

parser = argparse.ArgumentParser(description='匯出規格文件為 Word 檔')
parser.add_argument('--project', type=str, default='MagicSort', help='teardowns/ 底下的目標專案名稱')
args = parser.parse_args()

# 設定路徑
base_dir = r"c:\Users\akkuo\.gemini\antigravity\build_Game"
target_dir = os.path.join(base_dir, "teardowns", args.project)
output_dir = r"C:\Users\akkuo\Desktop\規格書"
os.makedirs(output_dir, exist_ok=True)

print(f"正在讀取 {target_dir} 的拆解規格，並將每個步驟獨立產出為 Word 檔...")

# 抓取所有 Markdown 檔案
md_files = glob.glob(os.path.join(target_dir, "*.md"))

for md_path in md_files:
    filename = os.path.basename(md_path)
    if "會議記錄" in filename: # 略過會議記錄，如果有需要可以註解掉這行
        continue
        
    doc = Document()
    doc.add_heading(f"{args.project} - {filename.replace('.md', '')} 規格書", 0)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line: 
            continue
        
        # 轉換 Markdown 基礎語法到 Word 樣式
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- ") or line.startswith("* "):
            # 處理粗體 **text**
            clean_line = line[2:].replace("**", "")
            doc.add_paragraph(clean_line, style='List Bullet')
        elif line.startswith(">"):
            p = doc.add_paragraph(line[1:].strip().replace("**", ""))
            p.style = 'Quote' # 若系統沒有Quote Style會回退為普通段落
        else:
            doc.add_paragraph(line.replace("**", ""))

    # 輸出檔名對應原始 markdown 檔名
    out_filename = f"{args.project}_{filename.replace('.md', '.docx')}"
    word_out_path = os.path.join(output_dir, out_filename)
    doc.save(word_out_path)
    print(f"[OK] 獨立產出 Word 成功: {word_out_path}")

print("Done! All exports completed.")
